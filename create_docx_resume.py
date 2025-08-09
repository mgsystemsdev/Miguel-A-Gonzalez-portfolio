#!/usr/bin/env python3
"""
Script to create a proper DOCX resume file using python-docx
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new Run object and add the text
    new_run = OxmlElement('w:r')

    # Set the run's style to hyperlink
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_job_entry(doc, title, company, dates, bullets):
    """Add a job entry with consistent formatting"""
    # Job title and company
    p = doc.add_paragraph()
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    
    p.add_run(f" | {company}")
    
    # Dates (right-aligned)
    date_p = doc.add_paragraph(dates)
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.space_after = Pt(6)
    
    # Bullet points
    for bullet in bullets:
        bullet_p = doc.add_paragraph(bullet, style='List Bullet')
        bullet_p.space_after = Pt(3)

def create_docx_resume():
    """Create a professional DOCX resume"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name and contact
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_run = header.add_run("Miguel A. Gonzalez Almonte")
    name_run.font.size = Pt(18)
    name_run.bold = True
    
    contact = doc.add_paragraph("Plano, TX • 787-367-9843 • mg.systems.dev@gmail.com")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    
    # Create paragraph with clickable links
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add hyperlinks
    add_hyperlink(links, "https://www.linkedin.com/in/miguel-gonzalez-8a389791", "LinkedIn")
    links.add_run(" • ")
    add_hyperlink(links, "https://mga210.github.io/DevProfile/", "Portfolio")
    links.add_run(" • ")
    add_hyperlink(links, "https://github.com/mga210", "GitHub")
    
    links.paragraph_format.space_after = Pt(12)
    
    # Professional tagline
    tagline = doc.add_paragraph("AI Systems Builder | GPT-Powered Workflow Architect | Operational Intelligence Technologist")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.runs[0]
    tagline_run.font.size = Pt(12)
    tagline_run.bold = True
    tagline.paragraph_format.space_after = Pt(12)
    
    # Professional Summary
    add_heading_with_style(doc, "PROFESSIONAL SUMMARY", 1)
    
    summary_text = """AI systems builder with deep roots in operations leadership and a self-driven path into software design, automation, and GPT agent architecture. I've built and deployed intelligent tools that replaced spreadsheets, reduced operational delays, and brought logic-based coordination into live operational environments. My systems are not experiments — they are active solutions used in production to enforce lifecycle integrity, audit state transitions, and support non-technical users.

Though I've never held a formal software title, I've architected AI assistants, Python dashboards, and decision frameworks that reflect enterprise-grade thinking and field-tested pragmatism. My work lives at the intersection of operational intelligence and systems clarity — where good logic can save time, prevent errors, and build trust.

Now seeking roles where I can continue designing agent-powered workflows, internal tools, and smart coordination systems — especially in environments that value practical intelligence, not just pedigree."""
    
    summary_para = doc.add_paragraph(summary_text)
    summary_para.paragraph_format.space_after = Pt(12)
    
    # Skills Summary
    add_heading_with_style(doc, "SKILLS SUMMARY", 1)
    
    skills_data = [
        ("AI & LLM Systems", [
            "GPT Agent Design, Prompt Architecture, Logic Scaffolding, 16-Block Canvas Systems",
            "Recursive Flows, Task Routing Systems, Lifecycle Modeling, Behavioral Frameworks",
            "Role-Based Interaction Flows, Educational AI Design, Cognitive Load Adaptation",
            "Clean Architecture Specification, GTPS Framework, Modular Dialogue Systems"
        ]),
        ("Python Development", [
            "Python 3 (Modular Scripting, DTO Structures, API Development, GUI Applications)",
            "Pandas (Data Structuring, Filtering, Export Pipelines, Analytics)",
            "PySide6 and Tkinter (GUI Architecture, Interactive Learning Environments, Desktop Applications)",
            "FastAPI (Lightweight API Services, Database Integration)",
            "SQLite and Supabase (Layered DB Use, State Tracing, Cloud Integration)"
        ]),
        ("Data & Dashboards", [
            "Excel (Advanced Formulas, VBA, Macros, Conditional Logic)",
            "Power BI (Custom Dashboards, Workflow Reporting, Analytics)",
            "CMMS Data Structuring, Delay Tracking Logic, Heatmaps and Export Systems",
            "Real-time Analytics, Service Manager Consoles, Performance Metrics"
        ]),
        ("No-Code / Low-Code Web Deployment", [
            "Replit: Web Page Design & Deployment (HTML Structuring, Visual Flow, Script Integration)",
            "Layout & Usability Architecture (Fonts, Spacing, Mobile-Friendliness, User Interaction)",
            "Integration of Tools: Formspree, Netlify, Cloudflare for Functional Delivery",
            "Full-stack launch using prebuilt components and code remixing",
            "UX-Focused Page Building: Professional design without hand-coding"
        ]),
        ("Workflow Automation", [
            "Legacy System Modernization, Excel-to-System Migration",
            "Process Optimization, Business Intelligence, Operations Intelligence",
            "Task Template Systems, Lifecycle Orchestration, State Management",
            "Rapid Web Application Development, Full-Stack Solutions"
        ])
    ]
    
    for category, skills in skills_data:
        cat_p = doc.add_paragraph()
        cat_run = cat_p.add_run(category)
        cat_run.bold = True
        cat_run.font.size = Pt(11)
        
        for skill in skills:
            skill_p = doc.add_paragraph(f"• {skill}")
            skill_p.paragraph_format.space_after = Pt(3)
    
    # Professional Experience
    add_heading_with_style(doc, "PROFESSIONAL EXPERIENCE", 1)
    
    # Job 1
    add_job_entry(doc, "Service Maintenance Manager", "MAA – Dallas, TX", "Jun 2023 – Present", [
        "Led service operations across 3 multifamily properties while designing and deploying the Make Ready Digital Board (DMRB) — a logic-based AI tool used live to coordinate unit readiness",
        "Replaced manual spreadsheets with a state-resolved Python system using DTOs, task templates, and lifecycle enforcement",
        "Built audit-safe, role-scoped task flows with offline queueing, conflict detection, and automatic readiness locking",
        "Reduced unit turnover time from 13–20 days to 7 through system-led coordination and real-time visibility",
        "Integrated Python (Pandas, PySide6, FastAPI), SQLite, and Supabase to enable full-stack functionality"
    ])
    
    # Job 2
    add_job_entry(doc, "Service Manager", "RPM Living – Dallas, TX", "May 2022 – Jun 2023", [
        "Directed daily service operations at a high-volume multifamily property, overseeing technician workflows, vendor schedules, and turnover timelines",
        "Built and deployed custom Excel dashboards to reduce admin friction, improve task tracking, and align team focus",
        "Applied Agile-style planning methods to improve technician coverage and reduce backlog",
        "Created SOPs and structured vendor workflows to streamline unit turnover"
    ])
    
    # Job 3
    add_job_entry(doc, "Operations Assistant → Kitchen Manager", "Universal Studios – Orlando, FL", "2009 – 2018", [
        "Progressed from operations assistant to leading two full-service kitchens during high-volume park operations",
        "Participated in engineering launch teams for new venues and attractions",
        "Led cross-kitchen menu rollout projects, coordinating timing, staff training, and guest flow readiness",
        "Developed early awareness of system bottlenecks, team handoffs, and operations logic under pressure"
    ])
    
    # Education & Certifications
    add_heading_with_style(doc, "EDUCATION & CERTIFICATIONS", 1)
    
    edu_p = doc.add_paragraph()
    edu_run = edu_p.add_run("Bachelor of Business Administration in Computer Information Systems")
    edu_run.bold = True
    edu_p.add_run("\nAna G. Méndez University – Carolina, PR (In Progress)")
    edu_p.paragraph_format.space_after = Pt(6)
    
    cert_p = doc.add_paragraph()
    cert_run = cert_p.add_run("Completed Certifications:")
    cert_run.bold = True
    
    certifications = [
        "Python for Everybody – University of Michigan / Coursera (2025)",
        "Python 3 – Intermediate Track (2025)",
        "Google Project Management Certificate – Coursera (2025)",
        "EPA Section 608 Certification – HVAC Systems (2018)"
    ]
    
    for cert in certifications:
        cert_bullet = doc.add_paragraph(f"• {cert}")
        cert_bullet.paragraph_format.space_after = Pt(3)
    
    # Key Projects
    add_heading_with_style(doc, "KEY PROJECTS", 1)
    
    projects = [
        ("Make Ready Digital Board (DMRB)", "AI-Powered Task Lifecycle System | Python, Pandas, PySide6, FastAPI, SQLite → Supabase", [
            "Replaced spreadsheets with a logic-resolved unit coordination engine deployed across 3 properties",
            "Reduced turnover from 13–20 days to 7 with role-based access, task gating, and offline queueing",
            "Features 13-screen interface including Service Manager Console, MR Board, Task Manager, and Analytics Dashboard",
            "Implemented core dialogs system for lifecycle orchestration and system-wide sync triggers"
        ]),
        ("Python Training Board (PTB)", "Interactive Learning Environment | Python, Tkinter, PySide6, GUI Development", [
            "Designed hands-on platform for GUI development training using sandbox-style experimentation",
            "Built modular demo launchers with authorization flow, framework selection hub, and training interfaces",
            "Features folder copy/isolate system, UI component explorer, and editable canvas with real-time code reflection",
            "Supports both PySide6 and Tkinter frameworks with structured layout training and code exploration mode"
        ]),
        ("System Pilot", "GPT-Powered Architecture Strategist | 16-Phase GTPS Framework, Clean Architecture", [
            "Designed deterministic GPT assistant operating through 3-phase pipeline with 16-phase GTPS framework",
            "Transforms software ideas into implementation-ready Clean Architecture blueprints with modular specifications",
            "Enforces architectural boundaries through structured file tree generation and module discovery interrogation",
            "Outputs YAML, JSON, MD formats with developer-ready specs and blocking protocol enforcement"
        ]),
        ("Blueprint Buddy", "Modular GPT Instruction Architect | 16-Block Logic Canvas, Validation Engines", [
            "Built full-scale GPT builder transforming user intent into logic-rich, exportable instruction sets",
            "Features 16-block logic canvas with embedded validation engines and multi-format output capabilities",
            "Performs intent-to-logic transformation with advanced prompt optimization and system logic optimization",
            "Educational AI system with mentor-like behavior and cognitive load adaptation for structured learning"
        ]),
        ("Meta Code Sensei", "Educational AI Assistant | Biological Metaphors, Mentor Behavior, Adaptive Learning", [
            "Developed sophisticated educational AI with biological metaphor framework for code instruction",
            "Implements mentor-like behavior with cognitive load adaptation and structured learning approaches",
            "Features 16-phase behavioral framework for personalized developer guidance and skill building",
            "Designed for beginner-stage developers using AI-first approach for rapid skill acquisition"
        ])
    ]
    
    for project_name, project_type, bullets in projects:
        proj_p = doc.add_paragraph()
        proj_name_run = proj_p.add_run(project_name)
        proj_name_run.bold = True
        proj_name_run.font.size = Pt(11)
        
        proj_type_p = doc.add_paragraph(project_type)
        proj_type_p.runs[0].italic = True
        proj_type_p.paragraph_format.space_after = Pt(3)
        
        for bullet in bullets:
            bullet_p = doc.add_paragraph(f"• {bullet}")
            bullet_p.paragraph_format.space_after = Pt(3)
    
    # Save the document
    doc.save('resume.docx')
    print("Professional DOCX resume created successfully: resume.docx")

if __name__ == "__main__":
    create_docx_resume()