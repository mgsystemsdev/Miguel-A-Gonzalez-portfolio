#!/usr/bin/env python3
"""
Resume Generator - Creates PDF and DOCX versions of Miguel's resume
with exact formatting and hyperlinks as specified
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import blue, black

def add_hyperlink(paragraph, url, text):
    """Add hyperlink to docx paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_docx_resume():
    """Create DOCX resume with exact formatting"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name and Contact Info
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_run = header.add_run("Miguel A. Gonzalez Almonte")
    name_run.font.size = Pt(18)
    name_run.bold = True
    
    header.add_run("\nPlano, TX • 787-367-9843 • mg.systems.dev@gmail.com\n")
    
    # Add hyperlinks for LinkedIn, Portfolio, GitHub
    linkedin_p = doc.add_paragraph()
    linkedin_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(linkedin_p, "https://linkedin.com/in/miguel-gonzalez-8a389791", "LinkedIn")
    linkedin_p.add_run(" • ")
    add_hyperlink(linkedin_p, "https://mga210.github.io/DevProfile", "Portfolio")
    linkedin_p.add_run(" • ")
    add_hyperlink(linkedin_p, "https://github.com/mgsystemsdev", "GitHub")
    
    doc.add_paragraph()  # Spacing
    
    # Target Roles
    target_roles = doc.add_paragraph()
    target_roles.add_run("Target Roles:").bold = True
    target_roles.add_run("\nAI Engineer • Internal Tools Developer • Workflow Automation Specialist • GPT Systems Architect")
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    summary_title = doc.add_paragraph()
    summary_title.add_run("AI Systems Builder | GPT Workflow Architect | Operational Intelligence Technologist").bold = True
    
    summary_text = doc.add_paragraph("""Systems-focused builder with deep roots in operations leadership and a forward trajectory in AI-powered automation. I design decision-support systems that replace spreadsheets with logic-driven tools — reducing delays, enforcing lifecycle logic, and supporting team coordination in real time.

My portfolio includes production-used agents, dashboards, and Python-based platforms built from scratch — not in theory, but in live operational environments. I thrive at the intersection of practical coordination and scalable AI tooling.""")
    
    doc.add_paragraph()  # Spacing
    
    # Core Tools & Technologies
    tech_title = doc.add_paragraph()
    tech_title.add_run("Core Tools & Technologies").bold = True
    
    doc.add_paragraph("Languages & Frameworks: Python • FastAPI • PySide6 • Tkinter • VBA • SQLite • Supabase")
    doc.add_paragraph("AI & LLMs: GPT-4 • LangChain • Prompt Engineering • Agent Workflows")
    doc.add_paragraph("Automation & Dashboards: Pandas • Power BI • CMMS Integration • Excel Macros")
    doc.add_paragraph("Deployment & UX: Replit • Netlify • HTML/CSS • Low-Code Architecture")
    doc.add_paragraph("DevOps & Logic Design: Git • Modular Architecture • DTOs • Lifecycle Mapping • Role-Based Flows")
    
    doc.add_paragraph()  # Spacing
    
    # Professional Experience
    exp_title = doc.add_paragraph()
    exp_title.add_run("Professional Experience").bold = True
    
    # Service Maintenance Manager
    job1_title = doc.add_paragraph()
    job1_title.add_run("Service Maintenance Manager — MAA – Dallas, TX").bold = True
    job1_title.add_run("\nJun 2023 – Present")
    
    job1_desc = doc.add_paragraph("""Led field ops while designing and deploying the Make Ready Digital Board (DMRB) — a production-grade Python system for turnover coordination

Eliminated spreadsheets with lifecycle logic, offline queuing, and task validation — reducing unit turnover from 13–20 days to 7 days

Delivered full-stack functionality with Python, SQLite, Supabase, and PySide6

Acted as both systems architect and end user — field-testing every feature in production""")
    
    # Service Manager
    job2_title = doc.add_paragraph()
    job2_title.add_run("Service Manager — RPM Living – Dallas, TX").bold = True
    job2_title.add_run("\nMay 2022 – Jun 2023")
    
    job2_desc = doc.add_paragraph("""Directed maintenance operations and designed Excel-based dashboards to streamline task management and vendor oversight

Introduced structured workflows and SOPs, leading to improved turnover timelines

Applied agile coordination frameworks to reduce backlog and optimize coverage""")
    
    # Additional positions
    positions = [
        ("Residential Renovation Lead (Contractor) — First Choice", "Jan 2021 – May 2022", 
         """Oversaw full-scope renovation projects with digital project tools and milestone tracking

Created protocols to reduce callbacks and improve delivery quality"""),
        
        ("Unit Upgrade Specialist (Contractor) — FSI", "Feb 2020 – Jan 2021",
         "Executed unit upgrades across multiple sites, implementing mobile-tracked workflows and trade sequencing improvements"),
        
        ("Maintenance Technician II — American Community", "Feb 2018 – Feb 2020",
         "Delivered multi-system repairs and supported CMMS implementation for better tracking and preventive logic"),
        
        ("Ops Assistant → Kitchen Manager — Universal Studios – Orlando, FL", "2009 – 2018",
         """Progressed to managing two high-volume kitchen teams; led coordination for new venue launches and menu rollouts

Developed early instincts for operational systems design under pressure""")
    ]
    
    for title, dates, desc in positions:
        job_title = doc.add_paragraph()
        job_title.add_run(title).bold = True
        job_title.add_run(f"\n{dates}")
        doc.add_paragraph(desc)
    
    doc.add_paragraph()  # Spacing
    
    # Education
    edu_title = doc.add_paragraph()
    edu_title.add_run("Education").bold = True
    
    edu_desc = doc.add_paragraph()
    edu_desc.add_run("Bachelor of Business Administration (Computer Information Systems)").bold = True
    edu_desc.add_run("\nAna G. Méndez University – In Progress\n• 40+ credits completed previously; currently resumed studies")
    
    doc.add_paragraph()  # Spacing
    
    # Certifications
    cert_title = doc.add_paragraph()
    cert_title.add_run("Certifications").bold = True
    
    completed_title = doc.add_paragraph()
    completed_title.add_run("Completed:").bold = True
    
    doc.add_paragraph("Google Project Management Certificate – Coursera, 2025")
    doc.add_paragraph("Python for Everybody – University of Michigan, 2025")
    doc.add_paragraph("Python 3: Intermediate Track – Coursera, 2025")
    doc.add_paragraph("EPA Section 608 Certification – HVAC Systems, 2018")
    
    progress_title = doc.add_paragraph()
    progress_title.add_run("In Progress:").bold = True
    
    doc.add_paragraph("IBM AI Engineering Certificate – Coursera")
    
    doc.add_paragraph()  # Spacing
    
    # Key Projects
    projects_title = doc.add_paragraph()
    projects_title.add_run("Key Projects").bold = True
    
    projects = [
        ("Make Ready Digital Board (DMRB)", "Python, Pandas, PySide6, FastAPI, SQLite → Supabase",
         """Lifecycle-based task manager for property teams; used in production

Reduced turnover time by 50% through role-based access and smart task gating"""),
        
        ("Blueprint Buddy", "GPT Instruction Builder | Prompt Engineering, Validation Logic",
         "Designed an agent to convert user input into deployable GPT instructions with multi-format output"),
        
        ("System Pilot", "GPT-Powered Architecture Strategist",
         "Modular assistant for converting raw product concepts into blueprint-level system plans"),
        
        ("Meta Code Sensei", "Python Learning GPT Agent",
         "Guided dev learners through 12-phase Python + architecture learning flows using GPT recursion"),
        
        ("Python Training Board (PTB)", "Tkinter + PySide6 GUI Learning Tool",
         "Designed for developers to practice full-stack design patterns in a controlled, modular environment")
    ]
    
    for project_name, tech, desc in projects:
        project_title = doc.add_paragraph()
        project_title.add_run(project_name).bold = True
        project_title.add_run(f"\n{tech}")
        doc.add_paragraph(desc)
    
    doc.add_paragraph()  # Spacing
    
    # Results & Metrics
    results_title = doc.add_paragraph()
    results_title.add_run("Results & Metrics").bold = True
    
    results_text = doc.add_paragraph("""• 65% process efficiency gain via AI agent deployment
• 13→7 days unit turnover through DMRB system
• 40% maintenance cost reduction via predictive logic
• $25K+ cost savings through Excel-to-Python transformation
• Trained 50+ team members on tools, logic, and execution workflows""")
    
    # Save document
    doc.save("resume.docx")
    print("DOCX resume created successfully!")

def create_pdf_resume():
    """Create PDF resume with exact formatting"""
    doc = SimpleDocTemplate("resume.pdf", pagesize=letter,
                          rightMargin=54, leftMargin=54,
                          topMargin=36, bottomMargin=36)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        alignment=1,  # Center
        textColor=black
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=11,
        alignment=1,  # Center
        spaceAfter=12
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6,
        textColor=black
    )
    
    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontSize=12,
        spaceBefore=6,
        spaceAfter=3,
        leftIndent=0
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leftIndent=0
    )
    
    story = []
    
    # Header
    story.append(Paragraph("Miguel A. Gonzalez Almonte", title_style))
    story.append(Paragraph("Plano, TX • 787-367-9843 • mg.systems.dev@gmail.com", contact_style))
    story.append(Paragraph('<a href="https://linkedin.com/in/miguel-gonzalez-8a389791" color="blue">LinkedIn</a> • <a href="https://mga210.github.io/DevProfile" color="blue">Portfolio</a> • <a href="https://github.com/mga210" color="blue">GitHub</a>', contact_style))
    
    story.append(Spacer(1, 12))
    
    # Target Roles
    story.append(Paragraph("<b>Target Roles:</b>", normal_style))
    story.append(Paragraph("AI Engineer • Internal Tools Developer • Workflow Automation Specialist • GPT Systems Architect", normal_style))
    
    story.append(Spacer(1, 12))
    
    # Professional Summary
    story.append(Paragraph("<b>AI Systems Builder | GPT Workflow Architect | Operational Intelligence Technologist</b>", normal_style))
    story.append(Paragraph("Systems-focused builder with deep roots in operations leadership and a forward trajectory in AI-powered automation. I design decision-support systems that replace spreadsheets with logic-driven tools — reducing delays, enforcing lifecycle logic, and supporting team coordination in real time.", normal_style))
    story.append(Paragraph("My portfolio includes production-used agents, dashboards, and Python-based platforms built from scratch — not in theory, but in live operational environments. I thrive at the intersection of practical coordination and scalable AI tooling.", normal_style))
    
    # Core Tools & Technologies
    story.append(Paragraph("Core Tools & Technologies", heading_style))
    story.append(Paragraph("Languages & Frameworks: Python • FastAPI • PySide6 • Tkinter • VBA • SQLite • Supabase", normal_style))
    story.append(Paragraph("AI & LLMs: GPT-4 • LangChain • Prompt Engineering • Agent Workflows", normal_style))
    story.append(Paragraph("Automation & Dashboards: Pandas • Power BI • CMMS Integration • Excel Macros", normal_style))
    story.append(Paragraph("Deployment & UX: Replit • Netlify • HTML/CSS • Low-Code Architecture", normal_style))
    story.append(Paragraph("DevOps & Logic Design: Git • Modular Architecture • DTOs • Lifecycle Mapping • Role-Based Flows", normal_style))
    
    # Professional Experience
    story.append(Paragraph("Professional Experience", heading_style))
    
    # Service Maintenance Manager
    story.append(Paragraph("<b>Service Maintenance Manager — MAA – Dallas, TX</b><br/>Jun 2023 – Present", job_title_style))
    story.append(Paragraph("Led field ops while designing and deploying the Make Ready Digital Board (DMRB) — a production-grade Python system for turnover coordination", normal_style))
    story.append(Paragraph("Eliminated spreadsheets with lifecycle logic, offline queuing, and task validation — reducing unit turnover from 13–20 days to 7 days", normal_style))
    story.append(Paragraph("Delivered full-stack functionality with Python, SQLite, Supabase, and PySide6", normal_style))
    story.append(Paragraph("Acted as both systems architect and end user — field-testing every feature in production", normal_style))
    
    # Service Manager
    story.append(Paragraph("<b>Service Manager — RPM Living – Dallas, TX</b><br/>May 2022 – Jun 2023", job_title_style))
    story.append(Paragraph("Directed maintenance operations and designed Excel-based dashboards to streamline task management and vendor oversight", normal_style))
    story.append(Paragraph("Introduced structured workflows and SOPs, leading to improved turnover timelines", normal_style))
    story.append(Paragraph("Applied agile coordination frameworks to reduce backlog and optimize coverage", normal_style))
    
    # Additional positions
    positions = [
        ("Residential Renovation Lead (Contractor) — First Choice", "Jan 2021 – May 2022", 
         ["Oversaw full-scope renovation projects with digital project tools and milestone tracking",
          "Created protocols to reduce callbacks and improve delivery quality"]),
        
        ("Unit Upgrade Specialist (Contractor) — FSI", "Feb 2020 – Jan 2021",
         ["Executed unit upgrades across multiple sites, implementing mobile-tracked workflows and trade sequencing improvements"]),
        
        ("Maintenance Technician II — American Community", "Feb 2018 – Feb 2020",
         ["Delivered multi-system repairs and supported CMMS implementation for better tracking and preventive logic"]),
        
        ("Ops Assistant → Kitchen Manager — Universal Studios – Orlando, FL", "2009 – 2018",
         ["Progressed to managing two high-volume kitchen teams; led coordination for new venue launches and menu rollouts",
          "Developed early instincts for operational systems design under pressure"])
    ]
    
    for title, dates, descriptions in positions:
        story.append(Paragraph(f"<b>{title}</b><br/>{dates}", job_title_style))
        for desc in descriptions:
            story.append(Paragraph(desc, normal_style))
    
    # Education
    story.append(Paragraph("Education", heading_style))
    story.append(Paragraph("<b>Bachelor of Business Administration (Computer Information Systems)</b><br/>Ana G. Méndez University – In Progress<br/>• 40+ credits completed previously; currently resumed studies", normal_style))
    
    # Certifications
    story.append(Paragraph("Certifications", heading_style))
    story.append(Paragraph("<b>Completed:</b>", normal_style))
    story.append(Paragraph("Google Project Management Certificate – Coursera, 2025", normal_style))
    story.append(Paragraph("Python for Everybody – University of Michigan, 2025", normal_style))
    story.append(Paragraph("Python 3: Intermediate Track – Coursera, 2025", normal_style))
    story.append(Paragraph("EPA Section 608 Certification – HVAC Systems, 2018", normal_style))
    
    story.append(Paragraph("<b>In Progress:</b>", normal_style))
    story.append(Paragraph("IBM AI Engineering Certificate – Coursera", normal_style))
    
    # Key Projects
    story.append(Paragraph("Key Projects", heading_style))
    
    projects = [
        ("Make Ready Digital Board (DMRB)", "Python, Pandas, PySide6, FastAPI, SQLite → Supabase",
         ["Lifecycle-based task manager for property teams; used in production",
          "Reduced turnover time by 50% through role-based access and smart task gating"]),
        
        ("Blueprint Buddy", "GPT Instruction Builder | Prompt Engineering, Validation Logic",
         ["Designed an agent to convert user input into deployable GPT instructions with multi-format output"]),
        
        ("System Pilot", "GPT-Powered Architecture Strategist",
         ["Modular assistant for converting raw product concepts into blueprint-level system plans"]),
        
        ("Meta Code Sensei", "Python Learning GPT Agent",
         ["Guided dev learners through 12-phase Python + architecture learning flows using GPT recursion"]),
        
        ("Python Training Board (PTB)", "Tkinter + PySide6 GUI Learning Tool",
         ["Designed for developers to practice full-stack design patterns in a controlled, modular environment"])
    ]
    
    for project_name, tech, descriptions in projects:
        story.append(Paragraph(f"<b>{project_name}</b><br/>{tech}", job_title_style))
        for desc in descriptions:
            story.append(Paragraph(desc, normal_style))
    
    # Results & Metrics
    story.append(Paragraph("Results & Metrics", heading_style))
    story.append(Paragraph("• 65% process efficiency gain via AI agent deployment", normal_style))
    story.append(Paragraph("• 13→7 days unit turnover through DMRB system", normal_style))
    story.append(Paragraph("• 40% maintenance cost reduction via predictive logic", normal_style))
    story.append(Paragraph("• $25K+ cost savings through Excel-to-Python transformation", normal_style))
    story.append(Paragraph("• Trained 50+ team members on tools, logic, and execution workflows", normal_style))
    
    # Build PDF
    doc.build(story)
    print("PDF resume created successfully!")

if __name__ == "__main__":
    print("Generating resume files...")
    create_docx_resume()
    create_pdf_resume()
    print("Both resume files generated successfully!")
    print("Files created: resume.docx and resume.pdf")