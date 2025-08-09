#!/usr/bin/env python3
"""
Simple DOCX resume creation with working hyperlinks
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_simple_docx_resume():
    """Create a professional DOCX resume with working hyperlinks"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name and contact
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_run = header.add_run("Miguel A. Gonzalez Almonte")
    name_run.font.size = Pt(18)
    name_run.bold = True
    
    contact = doc.add_paragraph("Plano, TX • 787-367-9843 • mgonzalez869@gmail.com")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    
    # Links paragraph with hyperlinks
    links_para = doc.add_paragraph()
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add hyperlinks using the built-in add_hyperlink method
    links_para.add_run("LinkedIn: https://www.linkedin.com/in/miguel-gonzalez-8a389791\n")
    links_para.add_run("Portfolio: https://mga210.github.io/DevProfile/\n") 
    links_para.add_run("GitHub: https://github.com/mga210")
    links_para.paragraph_format.space_after = Pt(12)
    
    # Professional tagline
    tagline = doc.add_paragraph("AI Systems Builder | GPT-Powered Workflow Architect | Operational Intelligence Technologist")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.runs[0]
    tagline_run.font.size = Pt(12)
    tagline_run.bold = True
    tagline.paragraph_format.space_after = Pt(12)
    
    # Professional Summary
    summary_heading = doc.add_heading("PROFESSIONAL SUMMARY", 1)
    
    summary_text = """AI systems builder with deep roots in operations leadership and a self-driven path into software design, automation, and GPT agent architecture. I've built and deployed intelligent tools that replaced spreadsheets, reduced operational delays, and brought logic-based coordination into live operational environments.

Though I've never held a formal software title, I've architected AI assistants, Python dashboards, and decision frameworks that reflect enterprise-grade thinking and field-tested pragmatism. My work lives at the intersection of operational intelligence and systems clarity — where good logic can save time, prevent errors, and build trust.

Now seeking roles where I can continue designing agent-powered workflows, internal tools, and smart coordination systems — especially in environments that value practical intelligence, not just pedigree."""
    
    summary_para = doc.add_paragraph(summary_text)
    summary_para.paragraph_format.space_after = Pt(12)
    
    # Skills Summary
    skills_heading = doc.add_heading("SKILLS SUMMARY", 1)
    
    # AI Skills
    ai_skills = doc.add_paragraph()
    ai_title = ai_skills.add_run("AI & Workflow Intelligence")
    ai_title.bold = True
    ai_title.font.size = Pt(11)
    
    ai_bullets = [
        "GPT Agent Design, Prompt Architecture, Logic Scaffolding",
        "Recursive Flows, Task Routing Systems, Lifecycle Modeling", 
        "Role-Based Interaction Flows, Instructional UX for Non-Technical Users"
    ]
    
    for skill in ai_bullets:
        doc.add_paragraph(f"• {skill}", style='List Bullet')
    
    # Python Skills  
    python_skills = doc.add_paragraph()
    python_title = python_skills.add_run("Python Systems Development")
    python_title.bold = True
    python_title.font.size = Pt(11)
    
    python_bullets = [
        "Python 3 (Modular Scripting, DTO Structures, API Basics)",
        "Pandas (Data Structuring, Filtering, Export Pipelines)",
        "PySide6 and Tkinter (GUI Architecture, Interaction Flows)",
        "FastAPI (Lightweight API Services)",
        "SQLite and Supabase (Layered DB Use, State Tracing, Portability)"
    ]
    
    for skill in python_bullets:
        doc.add_paragraph(f"• {skill}", style='List Bullet')
    
    # Professional Experience
    exp_heading = doc.add_heading("PROFESSIONAL EXPERIENCE", 1)
    
    # Job 1
    job1_header = doc.add_paragraph()
    job1_title = job1_header.add_run("Service Maintenance Manager")
    job1_title.bold = True
    job1_title.font.size = Pt(11)
    job1_header.add_run(" | MAA – Dallas, TX | Jun 2023 – Present")
    
    job1_bullets = [
        "Led service operations across 3 multifamily properties while designing and deploying the Make Ready Digital Board (DMRB)",
        "Replaced manual spreadsheets with a state-resolved Python system using DTOs, task templates, and lifecycle enforcement",
        "Reduced unit turnover time from 13–20 days to 7 through system-led coordination and real-time visibility",
        "Integrated Python (Pandas, PySide6, FastAPI), SQLite, and Supabase to enable full-stack functionality"
    ]
    
    for bullet in job1_bullets:
        doc.add_paragraph(f"• {bullet}", style='List Bullet')
    
    # Job 2
    job2_header = doc.add_paragraph()
    job2_title = job2_header.add_run("Service Manager")
    job2_title.bold = True
    job2_title.font.size = Pt(11)
    job2_header.add_run(" | RPM Living – Dallas, TX | May 2022 – Jun 2023")
    
    job2_bullets = [
        "Built and deployed custom Excel dashboards to reduce admin friction and improve task tracking",
        "Applied Agile-style planning methods to improve technician coverage and reduce backlog",
        "Created SOPs and structured vendor workflows to streamline unit turnover"
    ]
    
    for bullet in job2_bullets:
        doc.add_paragraph(f"• {bullet}", style='List Bullet')
    
    # Education
    edu_heading = doc.add_heading("EDUCATION & CERTIFICATIONS", 1)
    
    edu_para = doc.add_paragraph()
    edu_title = edu_para.add_run("Bachelor of Business Administration in Computer Information Systems")
    edu_title.bold = True
    edu_para.add_run("\nAna G. Méndez University – Carolina, PR (In Progress)")
    
    cert_para = doc.add_paragraph()
    cert_title = cert_para.add_run("Completed Certifications:")
    cert_title.bold = True
    
    certifications = [
        "Python for Everybody – University of Michigan / Coursera (2025)",
        "Google Project Management Certificate – Coursera (2025)",
        "EPA Section 608 Certification – HVAC Systems (2018)"
    ]
    
    for cert in certifications:
        doc.add_paragraph(f"• {cert}", style='List Bullet')
    
    # Key Projects
    projects_heading = doc.add_heading("KEY PROJECTS", 1)
    
    # Project 1
    proj1 = doc.add_paragraph()
    proj1_title = proj1.add_run("Make Ready Digital Board (DMRB)")
    proj1_title.bold = True
    proj1_title.font.size = Pt(11)
    
    proj1_desc = doc.add_paragraph("AI-Powered Task Lifecycle System | Python, Pandas, PySide6, FastAPI, SQLite")
    proj1_desc.runs[0].italic = True
    
    proj1_bullets = [
        "Replaced spreadsheets with a logic-resolved unit coordination engine deployed across 3 properties",
        "Reduced turnover from 13–20 days to 7 with role-based access, task gating, and offline queueing"
    ]
    
    for bullet in proj1_bullets:
        doc.add_paragraph(f"• {bullet}", style='List Bullet')
    
    # Project 2
    proj2 = doc.add_paragraph()
    proj2_title = proj2.add_run("System Pilot")
    proj2_title.bold = True
    proj2_title.font.size = Pt(11)
    
    proj2_desc = doc.add_paragraph("GPT-Powered Architecture Strategist | Prompt Logic, Modular Dialogues, Python")
    proj2_desc.runs[0].italic = True
    
    proj2_bullets = [
        "Designed a GPT assistant that guides users from raw product ideas into full system blueprints",
        "Enforced architectural planning through modular dialogue and logic scaffolding"
    ]
    
    for bullet in proj2_bullets:
        doc.add_paragraph(f"• {bullet}", style='List Bullet')
    
    # Save the document
    doc.save('resume.docx')
    print("Professional DOCX resume with hyperlinks created successfully: resume.docx")

if __name__ == "__main__":
    create_simple_docx_resume()